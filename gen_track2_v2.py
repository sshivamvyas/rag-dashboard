"""Generate Track 2 Revised Report (v2 with regex entity extraction)."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = r"C:\Users\adm\.openclaw-autoclaw\agents\task-2-rag\workspace\rag-study\DELIVERY\Track2_LegalBenchRAG_Revised_Report.docx"

doc = Document()

# ─── Styles ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Calibri')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    hs.font.name = 'Calibri'

def add_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def add_bold(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p

# ═══════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('RAG Engineering Comparative Study')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run('Track 2: Graph-Based Retrieval')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Revised Report — Regex Entity Extraction')
r.font.size = Pt(14)
r.font.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Dataset:  LegalBench-RAG (200 contracts, 53,044 chunks, 30 QA pairs)').font.size = Pt(11)
info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
info2.add_run('LLM:      google/flan-t5-xl (3B, local, $0)').font.size = Pt(11)
info3 = doc.add_paragraph()
info3.alignment = WD_ALIGN_PARAGRAPH.CENTER
info3.add_run('Embed:    all-MiniLM-L6-v2 (384d) | Graph: NetworkX entity co-occurrence').font.size = Pt(11)
info4 = doc.add_paragraph()
info4.alignment = WD_ALIGN_PARAGRAPH.CENTER
info4.add_run('Hardware: Kaggle T4 GPU (free tier) | Date: July 2026').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Pipeline Overview',
    '3. Baseline (Track 1)',
    '4. Stage 2 — Structured vs Fixed-Token Chunking',
    '5. Stage 3 — Entity Extraction (Regex-based)',
    '6. Stage 4 — Graph Construction + Local Search (1-hop)',
    '7. Stage 5 — Community Summarization + Global Search',
    '8. Stage 6 — Query Router',
    '9. Final Comparison: Track 2 vs Baseline',
    '10. Comparison: v1 vs v2 Entity Extraction',
    '11. Key Findings',
    '12. Root Cause Analysis',
    '13. Recommendations',
    '14. Limitations & Next Steps',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This report presents the revised results of Track 2 (Graph-Based Retrieval) for the '
    'RAG Engineering Comparative Study on the LegalBench-RAG dataset. The original Track 2 '
    'implementation used FLAN-T5-XL for entity extraction, producing only 16 unique entities '
    '(53% retrieval fallback rate). This revision replaces LLM-based entity extraction with '
    'regex-based extraction, yielding 6,659 unique entities and 0% fallback rate.'
)
doc.add_paragraph(
    'Despite this 416× improvement in graph structure, Track 2 still does not outperform the '
    'Track 1 baseline (dense retrieval). Graph-1hop retrieval achieves h@5=0.600 vs baseline h@5=0.667, '
    'and the overall routed system wins only 1/5 metrics (precision). This reveals a fundamental '
    'limitation: entity co-occurrence graphs create semantically narrow neighborhoods. When a chunk '
    'mentions "section 6." and "company X," its 1-hop neighbors all mention "section 6." — but from '
    'the same contractual context, reducing search diversity. The pipeline architecture is validated, '
    'but entity co-occurrence-based graph retrieval does not add value over dense search for this dataset.'
)

# ═══════════════════════════════════
# 2. PIPELINE OVERVIEW
# ═══════════════════════════════════
doc.add_heading('2. Pipeline Overview', level=1)
doc.add_paragraph(
    'Track 2 implements a 6-stage graph-based retrieval pipeline. All models and infrastructure '
    'are identical to Track 1 (FLAN-T5-XL for answer generation, all-MiniLM-L6-v2 for embeddings, '
    'Kaggle T4 GPU). The stages are:'
)
stages = [
    ('S2 — Structured Chunking', 'Compare fixed-token vs structured (section/paragraph) chunking strategies.'),
    ('S3 — Entity Extraction', 'Extract named entities (PARTY, CLAUSE, TERM, DATE, AMOUNT) from all chunks.'),
    ('S4 — Graph + Local Search', 'Build entity co-occurrence graph; retrieve via 1-hop neighborhood traversal.'),
    ('S5 — Community Search', 'Detect connected components → summarize communities → search via embedding similarity.'),
    ('S6 — Query Router', 'Classify query type (factual / relationship / synthesis) → route to appropriate search method.'),
]
for name, desc in stages:
    p = doc.add_paragraph()
    r = p.add_run(name + ': ')
    r.bold = True
    p.add_run(desc)

# ═══════════════════════════════════
# 3. BASELINE
# ═══════════════════════════════════
doc.add_heading('3. Baseline (Track 1)', level=1)
doc.add_paragraph(
    'The Track 1 baseline uses fixed-token chunking (256 tokens, 64 overlap) + all-MiniLM-L6-v2 '
    'dense retrieval + FLAN-T5-XL answer generation. These results serve as the reference for '
    'all Track 2 comparisons.'
)
add_table(doc,
    ['Metric', 'Baseline'],
    [['faithfulness', '0.617'],
     ['precision', '0.987'],
     ['recall', '0.803'],
     ['relevance', '0.733'],
     ['hit_rate@5', '0.667'],
     ['latency (ms)', '1,646']],
    'Table 1: Track 1 Baseline Metrics'
)

# ═══════════════════════════════════
# 4. S2 — STRUCTURED CHUNKING
# ═══════════════════════════════════
doc.add_heading('4. Stage 2 — Structured vs Fixed-Token Chunking', level=1)
doc.add_paragraph(
    'Structured chunking (StructuredChunker) splits legal text by section boundaries, paragraphs, '
    'and sub-paragraphs, producing 43,986 chunks. Fixed-token chunking (Chunker.fixed_token) uses '
    'a flat 256-token sliding window with 64-token overlap, producing 53,044 chunks.'
)
add_table(doc,
    ['Method', 'Chunks', 'Faith', 'Prec', 'Recall', 'Rel', 'H@5'],
    [['Fixed-Token', '53,044', '0.667', '0.987', '0.820', '0.653', '0.667'],
     ['Structured', '43,986', '0.433', '0.987', '0.837', '0.640', '0.767']],
    'Table 2: Structured vs Fixed-Token Chunking'
)
doc.add_paragraph(
    'Fixed-token chunking wins decisively. Structured chunking reduces faithfulness by -0.233 '
    '(0.667 → 0.433). The reason: section boundaries fragment cross-references between clauses. '
    'A single QA pair often requires evidence spanning multiple sections; structured chunking '
    'isolates each section into separate chunks, breaking the chain. Fixed-token carries forward.'
)

# ═══════════════════════════════════
# 5. S3 — ENTITY EXTRACTION
# ═══════════════════════════════════
doc.add_heading('5. Stage 3 — Entity Extraction (Regex-based)', level=1)
doc.add_paragraph(
    'Entity extraction uses five regex patterns operating on all 53,044 chunks. No LLM calls are '
    'used for extraction — regex is fast (<1 min for all chunks), free, and deterministic.'
)
ext_types = [
    ('PARTY', 'Capitalized multi-word + legal suffix (Inc./LLC/Corp/etc.)', '7,692'),
    ('CLAUSE', 'Section/Article/Clause + number references', '24,386'),
    ('TERM', 'Quoted defined terms (4-50 chars, capitalized)', '4,235'),
    ('DATE', 'Month Day, Year format', '2,631'),
    ('AMOUNT', 'Dollar amounts ($X million/billion/thousand)', '1,967'),
]
add_table(doc,
    ['Type', 'Pattern', 'Count'],
    ext_types,
    'Table 3: Entity Type Distribution (6,659 unique entities)'
)
doc.add_paragraph(
    'Total raw extractions: 52,170 from 53,044 chunks. After deduplication by (name, type) pairs, '
    '6,659 unique entities remain. The CLAUSE type dominates (24,386) because "Section X." patterns '
    'appear in nearly every chunk. Top entities by frequency: "section 6." (1,320), "section 5." (1,130), '
    '"section 4." (975) — highlighting the repetitive nature of section references in contracts.'
)
doc.add_paragraph(
    'Comparison with v1 (FLAN-T5): 179 raw entities, only 16 unique after dedup. The regex approach '
    'produces 416× more unique entities. However, the granularity differs: regex captures '
    'section-level entities which are highly repetitive (same entities across thousands of chunks), '
    'while FLAN-T5 captured abstract types ("right," "term") which had zero specificity. '
    'Regex wins on quantity and specificity.'
)

# ═══════════════════════════════════
# 6. S4 — GRAPH + LOCAL SEARCH
# ═══════════════════════════════════
doc.add_heading('6. Stage 4 — Graph Construction + Local Search (1-hop)', level=1)
doc.add_paragraph(
    'The entity co-occurrence graph is built by adding an edge between every pair of entities '
    'that appear in the same chunk. This produces a dense, connected graph with 6,656 nodes '
    'and 41,493 edges (avg degree 12.5).'
)
add_table(doc,
    ['Property', 'Value'],
    [['Nodes', '6,656'],
     ['Edges', '41,493'],
     ['Avg degree', '12.5'],
     ['Connected components', '1,809'],
     ['Top entity (degree)', "'section 3.' (408)"],
     ['Fallback rate', '0% (0/30 queries)']],
    'Table 4: Graph Properties'
)
doc.add_paragraph(
    'The graph-1hop search extracts entities from the question, traverses 1-hop neighbors, '
    'collects chunks from all found entities and their neighbors, then runs dense retrieval '
    'on the filtered chunk set.'
)
add_table(doc,
    ['Method', 'Faith', 'H@5', 'Prec', 'Recall', 'Rel'],
    [['Dense (ref)', '0.667', '0.667', '0.987', '0.820', '0.653'],
     ['Graph-1hop', '0.617', '0.600', '1.000', '0.753', '0.553']],
    'Table 5: Graph-1hop vs Dense Retrieval (All 30 QA)'
)
doc.add_paragraph(
    'Despite a dense graph with 0% fallback, graph-1hop underperforms on 4/5 metrics. Precision '
    'improves slightly (1.000 vs 0.987) — the filtered neighborhood contains fewer irrelevant chunks. '
    'But faithfulness, recall, relevance, and hit_rate all decline. The narrow neighborhood lacks '
    'diverse evidence needed for complete answers.'
)
add_table(doc,
    ['Method', 'Faith', 'H@5'],
    [['Dense', '0.667', '0.704'],
     ['Graph-1hop', '0.593', '0.630']],
    'Table 6: Relationship-Type Query Performance (27 QA pairs)'
)

# ═══════════════════════════════════
# 7. S5 — COMMUNITY SEARCH
# ═══════════════════════════════════
doc.add_heading('7. Stage 5 — Community Summarization + Global Search', level=1)
doc.add_paragraph(
    'Connected component analysis finds 1,809 components from the 6,656-node graph. Of these, '
    'only 1 component has ≥3 entities (the minimum threshold for community summarization). '
    'This community summarizes to "Parent\'s Tax Assets..." — a narrow topic covering only entity tax '
    'clauses.'
)
doc.add_paragraph(
    '26 synthesis-type queries are identified (keyword-based: "compare," "difference," "between," etc.). '
    'Community search is evaluated against dense retrieval for these queries.'
)
add_table(doc,
    ['Method', 'Faith', 'Recall', 'Rel'],
    [['Dense', '0.654', '0.827', '0.638'],
     ['Community', '0.623', '0.504', '0.527']],
    'Table 7: Community Search vs Dense (Synthesis Queries)'
)
doc.add_paragraph(
    'Community search is significantly worse: recall drops -0.323 and relevance drops -0.111. '
    'A single community cannot represent the diverse topics covered by 26 synthesis queries. '
    'Where dense retrieval searches the full corpus, community search is limited to one '
    'summary covering Parent Tax Assets — irrelevant to most questions.'
)

# ═══════════════════════════════════
# 8. S6 — QUERY ROUTER
# ═══════════════════════════════════
doc.add_heading('8. Stage 6 — Query Router', level=1)
doc.add_paragraph(
    'The query router classifies each question and routes to the appropriate search method: '
    'factual → dense, relationship → graph-1hop, synthesis → community search. The classifier '
    'uses keyword matching (relationship: "between," "parties," "relate"; synthesis: "compare," '
    '"difference," "versus," "summarize").'
)
doc.add_paragraph(
    'Routing distribution: relationship=27, factual=3, synthesis=0. The "between" keyword matches '
    '27/30 questions (90%), over-classifying nearly everything as "relationship-type." Synthesis '
    'keywords are subsumed by the relationship classifier due to keyword ordering. Fallback rate: '
    '0% — all 27 relationship queries route to graph-1hop and match entities.'
)
add_table(doc,
    ['Method', 'Faith', 'Prec', 'Recall', 'Rel', 'H@5'],
    [['Baseline (T1)', '0.617', '0.987', '0.803', '0.733', '0.667'],
     ['Routed (T2)', '0.600', '1.000', '0.753', '0.547', '0.600']],
    'Table 8: Routed (T2) vs Baseline (Final Comparison)'
)
add_table(doc,
    ['Route', 'n', 'Faith', 'H@5'],
    [['dense', '3', '0.667', '0.333'],
     ['graph-1hop', '27', '0.593', '0.630']],
    'Table 9: Per-Route Breakdown'
)
doc.add_paragraph(
    'The graph-1hop route produces worse results than dense on all metrics except precision. '
    'Notably, the dense route (only 3 factual queries) has lower h@5 (0.333) than graph-1hop (0.630) '
    'on relationship queries — but both are below the baseline dense score (0.667). '
    'The router adds overhead without improvement.'
)

# ═══════════════════════════════════
# 9. FINAL COMPARISON
# ═══════════════════════════════════
doc.add_heading('9. Final Comparison: Track 2 vs Baseline', level=1)
add_table(doc,
    ['Metric', 'Baseline', 'Track 2', 'Δ', 'Verdict'],
    [['faithfulness', '0.617', '0.600', '-0.017', 'BASELINE'],
     ['precision', '0.987', '1.000', '+0.013', 'TRACK 2'],
     ['recall', '0.803', '0.753', '-0.050', 'BASELINE'],
     ['relevance', '0.733', '0.547', '-0.186', 'BASELINE'],
     ['hit_rate@5', '0.667', '0.600', '-0.067', 'BASELINE'],
     ['latency (ms)', '1,646', '2,633', '+987', 'INFO']],
    'Table 10: Final Head-to-Head Comparison'
)
doc.add_paragraph(
    'Baseline wins 4/5 metrics. Track 2 wins only precision — the filtered graph neighborhood '
    'reduces noise but at the cost of completeness and relevance. The latency penalty (+987ms) '
    'comes from per-query index rebuilds in graph-1hop search.'
)

# ═══════════════════════════════════
# 10. v1 vs v2 COMPARISON
# ═══════════════════════════════════
doc.add_heading('10. Comparison: v1 (LLM) vs v2 (Regex) Entity Extraction', level=1)
doc.add_paragraph(
    'The change from LLM-based to regex-based entity extraction is the single largest improvement in '
    'Track 2. This section documents the transformation.'
)
add_table(doc,
    ['Property', 'v1 (FLAN-T5)', 'v2 (Regex)'],
    [['Unique entities', '16', '6,659'],
     ['Graph nodes', '16', '6,656'],
     ['Graph edges', '1', '41,493'],
     ['Fallback rate', '53%', '0%'],
     ['Entity quality', 'Generic types ("right", "party")', 'Specific names ("section 6.", "company X")'],
     ['Communities', '0', '1'],
     ['h@5 (graph-1hop)', '0.433', '0.600'],
     ['h@5 (routed)', '0.433', '0.600'],
     ['LLM calls', '203 (entity extraction)', '0 (entity extraction)']],
    'Table 11: v1 vs v2 Entity Extraction'
)
doc.add_paragraph(
    'The v2 regex approach produces a graph that is 416× larger in nodes and 41,493× larger in edges. '
    'Fallback drops from 53% to 0%. However, the h@5 improvement (0.433 → 0.600) is modest compared '
    'to the graph growth — confirming that the bottleneck shifted from graph construction to the '
    'fundamental retrieval strategy (1-hop traversal).'
)

# ═══════════════════════════════════
# 11. KEY FINDINGS
# ═══════════════════════════════════
doc.add_heading('11. Key Findings', level=1)

findings = [
    ('Regex extraction beats LLM: 6,659 vs 16 unique entities (416× improvement). FLAN-T5 cannot extract specific legal entity names; regex from legal text patterns is more effective and zero-cost.',
     'Table 11'),
    ('0% fallback rate. All 30 queries now match entities in the graph (was 53%). The graph is structurally complete.',
     'Section 6, Table 4'),
    ('Dense graph ≠ better retrieval. Graph-1hop loses to dense on 4/5 metrics despite 40k+ edges and 0% fallback.',
     'Section 6, Table 5'),
    ('Entity co-occurrence neighborhoods are semantically narrow. 1-hop neighbors all reference the same entity — from the same contractual context. Dense retrieval over the full corpus finds more diverse, complementary passages.',
     'Root Cause Analysis'),
    ('Only 1 community formable from 1,809 components. Most components are single entities; few have ≥3 entities for community summarization.',
     'Section 7'),
    ('Community search fails: -0.323 recall vs dense. A single community ("Parent\'s Tax Assets") cannot cover diverse synthesis queries.',
     'Section 7, Table 7'),
    ('Keyword-based query classification is weak: "between" matches 27/30 questions (90%). Synthesis keywords are subsumed, factual queries are rare.',
     'Section 8'),
    ('Structured chunking reduces faithfulness by -0.233 vs fixed-token. Section boundaries fragment cross-references between clauses.',
     'Section 4, Table 2'),
    ('The CLAUSE entity type dominates (24,386 / 52,170 raw extractions) because "Section X." patterns are highly repetitive. This floods the graph with generic nodes.',
     'Section 5, Table 3'),
    ('Track 2 adds +987ms latency vs baseline due to per-query index rebuild in graph-1hop — a 60% slowdown for worse results.',
     'Section 9, Table 10'),
]

for i, (finding, ref) in enumerate(findings, 1):
    p = doc.add_paragraph()
    r = p.add_run(f'{i}. {finding.split(".")[0]}.')
    r.bold = True
    p.add_run(f' {finding.split(".")[1]}.')
    p2 = doc.add_paragraph(f'   Evidence: {ref}')
    p2.runs[0].font.size = Pt(9)
    p2.runs[0].italic = True

# ═══════════════════════════════════
# 12. ROOT CAUSE ANALYSIS
# ═══════════════════════════════════
doc.add_heading('12. Root Cause Analysis', level=1)
doc.add_paragraph(
    'The fundamental failure mode of entity co-occurrence graphs for legal RAG is locality bias. '
    'Consider: chunk A mentions "Section 6.01" and "BridgeBio Pharma Inc." Chunk B mentions '
    '"Section 6.01" and "net operating loss." The entity co-occurrence graph adds an edge between '
    '"Section 6.01" and "BridgeBio Pharma Inc." and between "Section 6.01" and "net operating loss." '
    'When a query contains "Section 6.01," the 1-hop traversal retrieves chunks B (net operating loss) '
    'and potentially other chunks mentioning "Section 6.01" — but all from the same contractual section.'
)
doc.add_paragraph(
    'The neighborhood is semantically homogeneous because entities co-occur within legal sections. '
    'A section-level entity ("Section 6.01") appears in a fixed set of chunks; its 1-hop neighbors '
    'are other entities from the same section. Dense retrieval, on the other hand, searches the full '
    '53k-chunk corpus based on embedding similarity, which finds semantically relevant passages '
    'from different contracts, different sections — providing diverse evidence for the LLM.'
)
doc.add_paragraph(
    'Additionally, the CLAUSE entity type floods the graph: 24,386 of 52,170 (47%) extractions are '
    'section references like "section 6." and "section 5." These are legitimate entities, but they '
    'are highly repetitive across contracts. The same "section 5." entity appears in thousands of chunks '
    '(frequency 1,130), creating dense connections that don\'t meaningfully distinguish between topics.'
)
doc.add_paragraph(
    'The result: graph-1hop retrieves a narrower, less diverse chunk set than dense retrieval across '
    'all 53k chunks. Precision improves (+0.013) because noise is filtered, but every other metric '
    'declines. The graph architecture is validated — the pipeline correctly executes entity extraction '
    '→ graph construction → traversal → retrieval — but the traversal strategy (1-hop co-occurrence) '
    'is fundamentally suboptimal for legal text.'
)

# ═══════════════════════════════════
# 13. RECOMMENDATIONS
# ═══════════════════════════════════
doc.add_heading('13. Recommendations', level=1)

recs = [
    ('Hierarchical Entity Types',
     'Use only PARTY→PARTY and PARTY→CLAUSE edges; exclude CLAUSE→CLAUSE edges. Section references '
     'dominate the graph (47% of extractions) and don\'t contribute meaningful cross-entity relationships. '
     'Entity-typed filtering would reduce noise and improve neighborhood quality.'),
    ('Semantic Community Detection',
     'Replace connected-component analysis with density-based or spectral clustering on the embedding '
     'space. Components are too fine-grained (1,809 components from 6,656 nodes); semantic clusters '
     'would produce fewer, more coherent communities.'),
    ('Multi-hop with Diversity Penalty',
     'Use 2-hop or 3-hop traversal with a diversity penalty (MMR-style) to avoid retrieving '
     'semantically redundant chunks from the same neighborhood. This would counteract the locality '
     'bias while still leveraging graph structure.'),
    ('Relationship-Aware Entity Graphs',
     'Instead of co-occurrence, use a stronger LLM to extract explicit relationships between entities '
     '("BridgeBio Pharma Inc. is a party to Section 6.01"). This produces a directed, typed graph '
     'where traversal follows semantic paths, not statistical co-occurrence.'),
    ('Skip 1-hop Entirely; Use Graph for Query Expansion',
     'Instead of filtering the retrieval set via 1-hop, use entity matching from the question to '
     'expand the query with related entity names, then run dense retrieval on the full corpus. '
     'This preserves search diversity while incorporating graph knowledge.'),
]

for title, desc in recs:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}: ')
    r.bold = True
    p.add_run(desc)

# ═══════════════════════════════════
# 14. LIMITATIONS & NEXT STEPS
# ═══════════════════════════════════
doc.add_heading('14. Limitations & Next Steps', level=1)

doc.add_heading('Limitations', level=2)
lims = [
    'The QA pair set (30 queries) is small — statistical significance is limited.',
    'The dataset subset (200 of 714 contracts) may not represent the full LegalBench-RAG distribution.',
    'Keyword-based query classification is brittle; an LLM-based classifier would be more accurate.',
    'Regex entity extraction misses implicit entities and is limited to 5 pattern types.',
    'No parameter sweep for graph traversal depth was conducted (only 1-hop tested).',
    'FLAN-T5-XL (3B) is relatively small for answer generation; a 7B+ model might benefit more from graph-filtered context.',
]
for lim in lims:
    doc.add_paragraph(lim, style='List Bullet')

doc.add_heading('Next Steps', level=2)
next_steps = [
    'Implement hierarchical entity filtering and re-run S4-S6.',
    'Test MMR-style diversity penalty in graph-1hop retrieval.',
    'Use graph for query expansion (entity-based) instead of retrieval filtering.',
    'Scale to full 714-contract LegalBench-RAG dataset.',
    'Build Track 3 (Context & Memory Optimization) for complete 3-track comparison.',
]
for ns in next_steps:
    doc.add_paragraph(ns, style='List Bullet')

# ─── Save ───
doc.save(OUT)
print(f'Report saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
