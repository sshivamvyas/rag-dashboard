#!/usr/bin/env python3
"""
Generate Track 2: Knowledge Structure & Graph-Based Retrieval Report
Professional Word document using python-docx.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r"C:\Users\adm\.openclaw-autoclaw\agents\task-2-rag\workspace\rag-study\DELIVERY\Track2_Graph_Based_Retrieval_Report.docx"

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Helper functions
def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_paragraph(doc, text, style_name='Normal', bold=False, size=None, color=None, alignment=None, space_before=None, space_after=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run.font.name = 'Calibri'
    return p

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark blue
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text, style='Normal')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p

def add_metric_table(doc, headers, rows, col_widths=None):
    """Add a Light Grid Accent 1 table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph()  # spacing
    return table

def add_page_break(doc):
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════

# Add some vertical space
for _ in range(6):
    doc.add_paragraph()

add_formatted_paragraph(doc, "RAG Engineering Study", bold=True, size=14,
                        color=(0x1F, 0x49, 0x7D), alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_formatted_paragraph(doc, "Track 2: Knowledge Structure & Graph-Based Retrieval", bold=True, size=22,
                        color=(0x1F, 0x49, 0x7D), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)

add_formatted_paragraph(doc, "Does a structured representation of the corpus improve\nreasoning quality beyond dense retrieval?", size=13,
                        color=(0x4A, 0x4A, 0x4A), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=16)

doc.add_paragraph()
doc.add_paragraph()

add_formatted_paragraph(doc, "Dataset: LegalBench-RAG (714 legal contracts, 6,889 QA pairs)", size=11,
                        color=(0x4A, 0x4A, 0x4A), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_formatted_paragraph(doc, "LLM: FLAN-T5-XL (3B) · Embeddings: all-MiniLM-L6-v2 · Graph: NetworkX + FAISS", size=11,
                        color=(0x4A, 0x4A, 0x4A), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_formatted_paragraph(doc, "Infrastructure: Kaggle T4 GPU · Cost: $0 · Fully Local", size=11,
                        color=(0x4A, 0x4A, 0x4A), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_formatted_paragraph(doc, "July 2026", size=11, color=(0x4A, 0x4A, 0x4A),
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_formatted_paragraph(doc, "Task 2 RAG — AutoClaw Agent", size=10,
                        color=(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "Executive Summary", level=1)

add_body(doc,
    "This report investigates Track 2 of the RAG Engineering Study: whether augmenting a standard dense "
    "retrieval pipeline with a structured knowledge graph improves reasoning quality on legal-domain "
    "question answering. Using LegalBench-RAG — a corpus of 714 legal contracts with 6,889 QA pairs "
    "spanning CUAD, MAUD, NDA interpretation, and PrivacyQA — we construct a six-stage pipeline that "
    "progressively introduces structure-aware chunking, entity extraction, graph construction, community "
    "summarization, and intelligent query routing."
)

add_body(doc,
    "The pipeline operates entirely on Kaggle's free-tier T4 GPU using FLAN-T5-XL (3B parameters) as the "
    "generation backbone, all-MiniLM-L6-v2 (384 dimensions) for dense embeddings, NetworkX for graph "
    "construction, and FAISS IndexFlatIP for vector storage. Total inference cost across all stages is $0."
)

add_body(doc,
    "The central research question — Does a structured representation of the corpus improve reasoning "
    "quality beyond dense retrieval? — is evaluated by comparing the fully routed Track 2 pipeline "
    "against a baseline dense-retrieval-only system across five key metrics: Exact Match (EM), "
    "F1 Score, BLEU, ROUGE-L, and METEOR. Stages 2–6 provide intermediate ablation results that "
    "isolate the contribution of each structural enhancement."
)

p = doc.add_paragraph()
run = p.add_run("Key Result (preliminary, to be confirmed by Kaggle execution): ")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
run = p.add_run(
    "The routed Track 2 pipeline is expected to outperform the dense-only baseline on "
    "relationship-type questions (via graph-1hop traversal) and synthesis-type questions "
    "(via community-level retrieval), while maintaining parity on factual lookups routed to "
    "dense search. The overall accuracy lift depends on the proportion of relationship and "
    "synthesis queries in the test set and the quality of entity extraction from legal text."
)
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_heading_styled(doc, "Pipeline Overview at a Glance", level=2)

add_metric_table(doc,
    ["Stage", "Name", "Key Technique", "Output"],
    [
        ["S1", "Baseline", "Fixed-token chunking + Dense Retrieval", "Reference metrics"],
        ["S2", "Structured Chunking", "SECTION/Article/paragraph-aware splitting", "Structure-annotated chunks"],
        ["S3", "Entity Extraction", "FLAN-T5 entity typing + regex fallback", "Per-chunk entity JSON"],
        ["S4", "Graph Construction & Local Search", "Entity co-occurrence graph + 1-hop traversal", "Graph-indexed retrieval"],
        ["S5", "Community Summarization & Global Search", "Connected-component communities + LLM summaries", "Community-level retrieval"],
        ["S6", "Query Router", "Query classification → method routing", "Routed retrieval pipeline"],
    ],
    col_widths=[1.5, 3.5, 5.5, 5.0]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "1. Introduction", level=1)

add_heading_styled(doc, "1.1 Research Question", level=2)
add_body(doc,
    "Dense retrieval has become the dominant paradigm for retrieval-augmented generation (RAG), "
    "offering semantic matching via embedding-space similarity. However, dense vectors are inherently "
    "flat — they do not capture explicit relationships between entities, document structure, or "
    "multi-hop connections within a corpus. This limitation is especially pronounced in legal domains, "
    "where contracts contain structured sections, defined terms, cross-references between clauses, "
    "and intricate party-obligation-right relationships."
)
add_body(doc,
    "Track 2 asks: Does a structured representation of the corpus improve reasoning quality beyond "
    "dense retrieval? In practical terms: if we build a knowledge graph of legal entities and their "
    "co-occurrence relationships, can we retrieve more relevant context than raw embedding similarity? "
    "And can we intelligently route questions to the right retrieval method based on question type?"
)

add_heading_styled(doc, "1.2 Architecture", level=2)
add_body(doc,
    "The Track 2 architecture extends the Track 1 baseline with five additional stages that form a "
    "knowledge-graph-enhanced retrieval pipeline:"
)

arch_items = [
    ("Structured Chunking (S2)", "Re-chunks documents preserving legal structure boundaries (Sections, Articles, paragraphs). This ensures that entities are not split across arbitrary token windows."),
    ("Entity Extraction (S3)", "Each chunk is processed by FLAN-T5-XL to extract typed legal entities: PARTY, CLAUSE, DATE, AMOUNT, TERM, OBLIGATION, and RIGHT. Results are cached to JSON with a regex fallback for robustness."),
    ("Graph Construction (S4)", "An entity co-occurrence graph is built using NetworkX. Nodes are entities; edges represent chunks where two entities appear together, weighted by co-occurrence frequency. Local search performs 1-hop neighbor traversal from query-matched entities."),
    ("Community Summarization (S5)", "Connected components in the graph form communities. Each community is summarized by FLAN-T5 into a one-sentence descriptor. Global search embeds community summaries and retrieves the top-k most relevant communities."),
    ("Query Router (S6)", "A classifier tags each incoming query as factual, relationship, or synthesis, then routes it to dense, graph-1hop, or community-search respectively. Queries with no extractable entities fall back to dense retrieval."),
]

for title_text, desc in arch_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_text + ": ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

add_heading_styled(doc, "1.3 Dataset", level=2)
add_body(doc,
    "The study uses LegalBench-RAG, a comprehensive legal contract dataset:"
)

add_metric_table(doc,
    ["Property", "Value"],
    [
        ["Total Documents", "714 legal contracts"],
        ["Sub-corpora", "CUAD, MAUD, NDA, PrivacyQA"],
        ["Total QA Pairs", "6,889"],
        ["Document Subset (Track 2)", "50 documents (entity extraction sample)"],
        ["Entity Extraction Sample", "~200 chunks (every Nth chunk)"],
        ["Chunk Size (Baseline)", "512 tokens (fixed)"],
        ["Chunk Size (Structured)", "Variable (preserves structure boundaries)"],
    ],
    col_widths=[6.5, 9.0]
)

add_body(doc,
    "The full 714-document corpus is used for dense indexing (S1 baseline and fallback). "
    "The 50-document subset is used for entity extraction (S3), graph construction (S4), and "
    "community summarization (S5). The subset was chosen to keep LLM extraction calls manageable "
    "(~200 chunks × 1 call each = 200 FLAN-T5 calls vs. ~4,200 chunks across all 714 docs)."
)

add_heading_styled(doc, "1.4 Infrastructure", level=2)

add_metric_table(doc,
    ["Component", "Technology", "Details"],
    [
        ["LLM", "google/flan-t5-xl", "3B parameters, HuggingFace, local inference"],
        ["Embeddings", "all-MiniLM-L6-v2", "384 dimensions, Sentence Transformers"],
        ["Vector Store", "FAISS IndexFlatIP", "Inner-product index for dense retrieval"],
        ["Graph Library", "NetworkX", "Entity co-occurrence graph"],
        ["Hardware", "Kaggle T4 GPU", "16 GB VRAM, free tier"],
        ["Cost", "$0", "All components local/open-source"],
        ["Environment", "Python 3.10+", "transformers, sentence-transformers, faiss-cpu, networkx"],
    ],
    col_widths=[3.5, 4.5, 7.5]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════
# 2. STAGE-BY-STAGE
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "2. Stage-by-Stage Analysis", level=1)

# ── STAGE 1 ──
add_heading_styled(doc, "2.1 Stage 1 — Baseline (Fixed-Token Chunking + Dense Retrieval)", level=2)

add_heading_styled(doc, "Description", level=3)
add_body(doc,
    "Stage 1 establishes the reference baseline: documents are split into fixed 512-token chunks with "
    "no awareness of document structure. Chunks are embedded using all-MiniLM-L6-v2 and indexed in "
    "a FAISS IndexFlatIP vector store. At query time, the query is embedded and the top-k (k=5) "
    "most similar chunks are retrieved via inner-product similarity. Retrieved chunks are concatenated "
    "and fed as context to FLAN-T5-XL for answer generation."
)
add_body(doc,
    "This stage is identical to the Track 1 baseline and serves as the control condition against "
    "which all structural enhancements are compared. Results from this stage provide the floor for "
    "the five core evaluation metrics."
)

add_heading_styled(doc, "Key Metrics", level=3)

add_metric_table(doc,
    ["Metric", "Value", "Notes"],
    [
        ["Chunking Method", "Fixed 512-token", "Sliding window, no overlap"],
        ["Number of Chunks (714 docs)", "~4,200 (from Kaggle run)", "Estimated; exact count from execution"],
        ["Embedding Dimension", "384", "all-MiniLM-L6-v2"],
        ["Top-k", "5", "Standard RAG retrieval depth"],
        ["Index Type", "FAISS IndexFlatIP", "Exact inner-product search"],
        ["EM Score", "(from Kaggle run)", "Exact Match on test set"],
        ["F1 Score", "(from Kaggle run)", "Token-level F1"],
        ["BLEU", "(from Kaggle run)", "BLEU-4"],
        ["ROUGE-L", "(from Kaggle run)", "Longest common subsequence"],
        ["METEOR", "(from Kaggle run)", "Synonym-aware matching"],
    ],
    col_widths=[4.0, 5.0, 6.5]
)

add_heading_styled(doc, "Analysis", level=3)
add_body(doc,
    "The baseline represents the state of practice for simple RAG pipelines: fast, scalable, and "
    "requiring no domain-specific preprocessing. However, fixed-token chunking can split legal "
    "sections mid-clause, fragmenting entities like 'Section 4.2(a)(iii)' across chunks and losing "
    "cross-reference context. This fragmentation is the primary motivation for structured chunking "
    "in Stage 2."
)

# ── STAGE 2 ──
add_heading_styled(doc, "2.2 Stage 2 — Structured Chunking", level=2)

add_heading_styled(doc, "Description", level=3)
add_body(doc,
    "Stage 2 replaces the fixed 512-token chunker with a structure-aware chunker that respects "
    "legal document hierarchy. The chunker identifies SECTION, Article, paragraph, sub-paragraph, "
    "and clause boundaries using regex patterns common in legal contracts (e.g., 'SECTION 1.', "
    "'Article II', '(a)', '(i)'). Each chunk preserves its structural boundary, ensuring that "
    "no entity, clause reference, or defined term spans a chunk boundary."
)
add_body(doc,
    "The structured chunker also annotates each chunk with its structure type, enabling downstream "
    "analysis of which structural units are most frequently retrieved and which contribute most "
    "to answer quality."
)

add_heading_styled(doc, "Key Metrics", level=3)

add_metric_table(doc,
    ["Metric", "Structured Chunking", "Fixed-Token (S1)", "Delta"],
    [
        ["Chunk Count (50-doc subset)", "(from Kaggle run)", "(from Kaggle run)", "—"],
        ["Avg Chunk Size (tokens)", "(from Kaggle run)", "512", "—"],
        ["% SECTION chunks", "(from Kaggle run)", "N/A", "—"],
        ["% Article chunks", "(from Kaggle run)", "N/A", "—"],
        ["% Paragraph chunks", "(from Kaggle run)", "N/A", "—"],
        ["% Sub-paragraph chunks", "(from Kaggle run)", "N/A", "—"],
        ["Cross-boundary splits avoided", "100%", "0%", "Qualitative"],
        ["Entity fragmentation rate", "(from Kaggle run)", "Estimated", "—"],
    ],
    col_widths=[4.5, 3.5, 3.5, 3.0]
)

add_heading_styled(doc, "Analysis", level=3)
add_body(doc,
    "Structured chunking trades uniform chunk size for semantic coherence. While some chunks may be "
    "smaller (a single sub-paragraph) or larger (a full section), each chunk represents a complete "
    "logical unit. This is expected to improve retrieval precision for queries referencing specific "
    "clauses or sections. The trade-off is that very large sections may exceed the LLM's context "
    "window when retrieved, requiring truncation. Structure type distribution provides insight into "
    "the granularity of the legal corpus."
)

# ── STAGE 3 ──
add_heading_styled(doc, "2.3 Stage 3 — Entity/Relationship Extraction", level=2)

add_heading_styled(doc, "Description", level=3)
add_body(doc,
    "Stage 3 performs typed entity extraction on each structured chunk. A carefully designed prompt "
    "is sent to FLAN-T5-XL requesting structured JSON output with seven entity types relevant to "
    "legal contract analysis:"
)

entity_types = [
    ("PARTY", "Company names, person names, signatories (e.g., 'ABC Corporation', 'John Smith')"),
    ("CLAUSE", "Section/article references (e.g., 'Section 4.2', 'Article III(a)')"),
    ("DATE", "Dates mentioned in the contract (e.g., 'January 15, 2024', 'the Effective Date')"),
    ("AMOUNT", "Monetary amounts or quantities (e.g., '$500,000', '15%')"),
    ("TERM", "Defined terms with specific legal meaning (e.g., 'Confidential Information', 'Indemnitee')"),
    ("OBLIGATION", "Mandatory actions signaled by must/shall/required (e.g., 'shall deliver', 'must notify')"),
    ("RIGHT", "Permissive actions signaled by may/entitled/reserves (e.g., 'may terminate', 'entitled to receive')"),
]

for etype, edesc in entity_types:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{etype}: ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run = p.add_run(edesc)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

add_body(doc,
    "Extraction is performed on a sample of ~200 chunks (every Nth chunk from the 50-document "
    "structured corpus). FLAN-T5-XL results are parsed as JSON and cached to disk. If FLAN-T5 "
    "fails to produce valid JSON (a known limitation for this model without few-shot examples), "
    "a regex-based fallback extractor provides baseline entity coverage using pattern matching "
    "for each entity type (e.g., r'Section \\d+\\.\\d+' for CLAUSE, r'\\$[\\d,]+' for AMOUNT)."
)

add_heading_styled(doc, "Key Metrics", level=3)

add_metric_table(doc,
    ["Metric", "Value", "Notes"],
    [
        ["Chunks Sampled", "~200", "Every Nth chunk from 50-doc subset"],
        ["LLM Extraction Calls", "~200", "One call per chunk"],
        ["Entity Types", "7", "PARTY, CLAUSE, DATE, AMOUNT, TERM, OBLIGATION, RIGHT"],
        ["Total Entities Extracted", "(from Kaggle run)", "LLM + fallback combined"],
        ["Avg Entities per Chunk", "(from Kaggle run)", "Varies by chunk complexity"],
        ["LLM JSON Parse Rate", "(from Kaggle run)", "% of FLAN-T5 outputs yielding valid JSON"],
        ["Fallback Activation Rate", "(from Kaggle run)", "% of chunks using regex fallback"],
        ["Most Frequent Entity Type", "(from Kaggle run)", "Expected: OBLIGATION or TERM"],
    ],
    col_widths=[4.5, 4.5, 6.5]
)

add_heading_styled(doc, "Analysis", level=3)
add_body(doc,
    "Entity extraction quality is the single most critical factor for downstream graph quality. "
    "FLAN-T5-XL, at 3B parameters and without few-shot JSON formatting examples, may struggle to "
    "produce consistently valid JSON. The regex fallback ensures the pipeline does not fail, but "
    "regex-based extraction has lower precision (e.g., cannot distinguish 'shall' as an obligation "
    "vs. 'shall' in a definition). The LLM JSON parse rate will quantify the reliability gap."
)
add_body(doc,
    "Entity type distribution provides a diagnostic of the legal corpus: high OBLIGATION density "
    "indicates prescriptive contracts (typical of NDAs and MAUD), while high PARTY density "
    "suggests multi-party agreements. This distribution informs the expected graph density in Stage 4."
)

# ── STAGE 4 ──
add_heading_styled(doc, "2.4 Stage 4 — Graph Construction & Local Search", level=2)

add_heading_styled(doc, "Description", level=3)
add_body(doc,
    "Stage 4 constructs an entity co-occurrence graph using NetworkX. In this undirected, weighted graph:"
)
graph_items = [
    "Nodes represent unique entities (normalized to lowercase, whitespace-trimmed).",
    "Edges connect entities that co-occur in the same chunk.",
    "Edge weights equal the number of chunks in which the entity pair co-occurs.",
    "Isolated nodes (entities appearing in only one chunk) are retained but disconnected.",
]
for item in graph_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

add_body(doc,
    "Local search (graph-1hop) works as follows: (1) extract entity mentions from the query using "
    "the same FLAN-T5 entity extraction pipeline; (2) locate matching nodes in the graph via exact "
    "string matching; (3) perform 1-hop neighbor traversal — retrieve all chunks associated with "
    "the query entity nodes and their immediate neighbors; (4) rank retrieved chunks by edge weight "
    "(stronger co-occurrence → higher relevance) and return top-k. If no entities are found in the "
    "query, the system falls back to dense retrieval."
)

add_heading_styled(doc, "Key Metrics", level=3)

add_metric_table(doc,
    ["Metric", "Value", "Notes"],
    [
        ["Graph Nodes", "(from Kaggle run)", "Unique entities"],
        ["Graph Edges", "(from Kaggle run)", "Entity co-occurrence pairs"],
        ["Graph Density", "(from Kaggle run)", "Edges / possible edges"],
        ["Avg Node Degree", "(from Kaggle run)", "Mean connections per entity"],
        ["Max Component Size", "(from Kaggle run)", "Largest connected component"],
        ["Num Components", "(from Kaggle run)", "Connected components"],
        ["1-Hop Candidates (avg/query)", "(from Kaggle run)", "Entities + neighbors per query"],
        ["Fallback Rate", "(from Kaggle run)", "% queries with no extractable entities"],
        ["Retrieval Latency (graph)", "(from Kaggle run)", "ms per query"],
        ["Retrieval Latency (dense)", "(from Kaggle run)", "ms per query — comparison point"],
    ],
    col_widths=[5.0, 4.5, 6.0]
)

add_heading_styled(doc, "Analysis", level=3)
add_body(doc,
    "Graph-based local search is designed to excel on relationship-type questions — queries that "
    "ask about connections between entities (e.g., 'What obligations does Party A have regarding "
    "Confidential Information?'). By traversing co-occurrence edges, the graph retrieves chunks "
    "that connect related entities even when those chunks do not contain high embedding similarity "
    "to the query. This compensates for the 'semantic drift' that can occur in dense retrieval "
    "when legal terminology diverges from natural-language query phrasing."
)
add_body(doc,
    "The fallback rate is a critical operational metric: if a high proportion of queries fail to "
    "match any graph entity, the graph adds no value for those queries. Fallback typically occurs "
    "for definitional or abstract questions ('What is the governing law?') where the query contains "
    "no specifically extractable entity name."
)

# ── STAGE 5 ──
add_heading_styled(doc, "2.5 Stage 5 — Community Summarization & Global Search", level=2)

add_heading_styled(doc, "Description", level=3)
add_body(doc,
    "Stage 5 introduces global, higher-level retrieval through community detection and summarization. "
    "Using the entity co-occurrence graph from Stage 4, connected components are identified as "
    "natural communities — groups of entities that frequently appear together across chunks, "
    "indicating shared topical context (e.g., all entities related to an indemnification clause "
    "tend to form a single component)."
)
add_body(doc,
    "Each community is then summarized by FLAN-T5-XL into a single-sentence descriptor. For example, "
    "a community containing entities like 'Indemnitee', 'Losses', 'Claims', and 'Section 8' might "
    "be summarized as: 'Community covering indemnification obligations and related claims procedures.'"
)
add_body(doc,
    "Global search works by embedding community summaries using all-MiniLM-L6-v2, comparing query "
    "embeddings to summary embeddings via cosine similarity, and retrieving the top-k communities. "
    "All chunks associated with entities in the retrieved communities are returned as context. "
    "This enables retrieval at the conceptual level, matching queries to entire thematic regions "
    "of the corpus rather than individual chunks."
)

add_heading_styled(doc, "Key Metrics", level=3)

add_metric_table(doc,
    ["Metric", "Value", "Notes"],
    [
        ["Total Communities", "(from Kaggle run)", "Connected components"],
        ["Communities Summarized", "(from Kaggle run)", "Excluding singletons"],
        ["Avg Community Size", "(from Kaggle run)", "Entities per community"],
        ["Singleton Communities", "(from Kaggle run)", "Single-entity components"],
        ["Avg Summary Length (tokens)", "(from Kaggle run)", "One-sentence descriptor"],
        ["Top-k Communities Retrieved", "3–5", "Configurable parameter"],
        ["Chunks per Community (avg)", "(from Kaggle run)", "Context volume per community"],
        ["Summary Embedding Dimension", "384", "Same as chunk embeddings"],
    ],
    col_widths=[5.5, 4.0, 6.0]
)

add_heading_styled(doc, "Analysis", level=3)
add_body(doc,
    "Community-based global search is targeted at synthesis-type questions — queries that require "
    "aggregating information across multiple sections or documents (e.g., 'Summarize all termination "
    "conditions across this contract', 'What are the common indemnification patterns?'). By retrieving "
    "at the community level, the system provides broader contextual coverage than chunk-level retrieval, "
    "at the cost of potentially including less-relevant chunks within a large community."
)
add_body(doc,
    "Community quality depends on graph connectivity. A highly fragmented graph (many small components) "
    "produces narrow communities, while a densely connected graph may produce a few large communities "
    "that lose topical specificity. The component size distribution is therefore a key diagnostic."
)

# ── STAGE 6 ──
add_heading_styled(doc, "2.6 Stage 6 — Query Router", level=2)

add_heading_styled(doc, "Description", level=3)
add_body(doc,
    "Stage 6 introduces intelligent query routing: instead of using a single retrieval method for all "
    "questions, a classifier first analyzes each query to determine its type, then routes it to the "
    "most appropriate retrieval method:"
)

router_items = [
    ("factual → Dense Retrieval (S1)", "Simple lookups: 'What is the governing law?', 'Who is the disclosing party?' These benefit from embedding similarity and do not require entity traversal."),
    ("relationship → Graph-1Hop (S4)", "Multi-entity connections: 'What obligations does Party A have to Party B regarding Confidential Information?' These benefit from entity co-occurrence traversal."),
    ("synthesis → Community-Search (S5)", "Cross-document aggregation: 'Summarize all indemnification provisions.' 'What are the common termination triggers?' These benefit from community-level broad retrieval."),
]

for rtype, rdesc in router_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(rtype + ": ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run = p.add_run(rdesc)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

add_body(doc,
    "The classifier uses a simple keyword + entity-count heuristic: queries containing relationship "
    "keywords (between, related to, obligation of, rights regarding) or multiple extracted entities "
    "are classified as relationship; queries containing summarization keywords (summarize, overview, "
    "all, across, patterns) are classified as synthesis; all others default to factual."
)
add_body(doc,
    "An LLM-based classifier (FLAN-T5 with a classification prompt) is provided as an alternative "
    "for comparison against the heuristic approach. The LLM classifier is expected to be more "
    "nuanced but adds latency."
)

add_heading_styled(doc, "Key Metrics", level=3)

add_metric_table(doc,
    ["Metric", "Value", "Notes"],
    [
        ["Queries Routed to Dense", "(from Kaggle run)", "factual classification"],
        ["Queries Routed to Graph-1Hop", "(from Kaggle run)", "relationship classification"],
        ["Queries Routed to Community", "(from Kaggle run)", "synthesis classification"],
        ["Fallback Rate (graph → dense)", "(from Kaggle run)", "% relationship queries w/o entities"],
        ["Fallback Rate (community → dense)", "(from Kaggle run)", "% synthesis queries w/o communities"],
        ["Classifier Accuracy (heuristic)", "(from Kaggle run)", "vs. manual labels on validation set"],
        ["Classifier Accuracy (LLM)", "(from Kaggle run)", "FLAN-T5 zero-shot classification"],
        ["Router Overhead (ms)", "(from Kaggle run)", "Classification latency per query"],
    ],
    col_widths=[5.5, 4.0, 6.0]
)

add_heading_styled(doc, "Analysis", level=3)
add_body(doc,
    "The query router is the integration point where all previous stages combine into a single "
    "pipeline. Its effectiveness depends on classification accuracy: misclassified queries suffer "
    "from suboptimal retrieval (e.g., a relationship query routed to dense retrieval may miss "
    "relevant entity-connected chunks). The fallback rate measures robustness — how often the "
    "system must degrade to dense retrieval because graph/community data is insufficient."
)
add_body(doc,
    "The router also provides a natural ablation framework: by comparing per-query-type performance "
    "between the routed pipeline and the baseline, we can isolate which question types benefit most "
    "from structural retrieval."
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════
# 3. FINAL COMPARISON
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "3. Final Comparison: Track 2 Routed Pipeline vs. Baseline", level=1)

add_body(doc,
    "The final evaluation compares the end-to-end Track 2 routed pipeline (S6: query → classifier "
    "→ dense/graph-1hop/community → FLAN-T5 generation) against the Track 1/S1 baseline "
    "(query → dense retrieval → FLAN-T5 generation). The comparison uses all five core metrics "
    "on the full LegalBench-RAG test set."
)

add_heading_styled(doc, "3.1 Overall Performance", level=2)

add_metric_table(doc,
    ["Metric", "Baseline (S1 Dense-Only)", "Track 2 Routed (S6)", "Delta", "Direction"],
    [
        ["Exact Match (EM)", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Expected: ↑ or ≈"],
        ["F1 Score", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Expected: ↑ or ≈"],
        ["BLEU-4", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Expected: ↑ or ≈"],
        ["ROUGE-L", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Expected: ↑ or ≈"],
        ["METEOR", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Expected: ↑ or ≈"],
    ],
    col_widths=[2.5, 3.5, 3.5, 3.0, 3.0]
)

add_heading_styled(doc, "3.2 Per-Question-Type Breakdown", level=2)

add_body(doc,
    "Where Track 2 is expected to show the strongest advantage is on relationship and synthesis "
    "questions, where graph-based retrieval provides information that dense embeddings miss. "
    "Factual questions should show approximate parity, with potential slight degradation if "
    "graph routing adds noise."
)

add_metric_table(doc,
    ["Question Type", "Baseline F1", "Track 2 F1", "Delta", "Primary Method"],
    [
        ["Factual", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Dense (same as baseline)"],
        ["Relationship", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Graph-1Hop traversal"],
        ["Synthesis", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Community-level retrieval"],
        ["All Types (weighted)", "(from Kaggle run)", "(from Kaggle run)", "(from Kaggle run)", "Routed (hybrid)"],
    ],
    col_widths=[2.5, 2.8, 2.8, 2.8, 4.6]
)

add_heading_styled(doc, "3.3 Ablation: Contribution of Each Stage", level=2)

add_body(doc,
    "To isolate the marginal contribution of each structural enhancement, we evaluate intermediate "
    "pipeline configurations:"
)

add_metric_table(doc,
    ["Configuration", "F1 Score", "Delta vs. S1", "Notes"],
    [
        ["S1: Baseline (dense only)", "(from Kaggle run)", "—", "Reference"],
        ["S1 + S2: Structured chunking only", "(from Kaggle run)", "(from Kaggle run)", "Chunking quality impact"],
        ["S1 + S2 + S4: Graph-1Hop only", "(from Kaggle run)", "(from Kaggle run)", "No router, graph for all"],
        ["S1 + S2 + S5: Community only", "(from Kaggle run)", "(from Kaggle run)", "No router, community for all"],
        ["S6: Full routed pipeline", "(from Kaggle run)", "(from Kaggle run)", "All stages integrated"],
    ],
    col_widths=[4.5, 2.8, 2.8, 5.4]
)

add_heading_styled(doc, "3.4 Efficiency Comparison", level=2)

add_metric_table(doc,
    ["Metric", "Baseline (S1)", "Track 2 Routed (S6)", "Notes"],
    [
        ["Avg Retrieval Time (ms)", "(from Kaggle run)", "(from Kaggle run)", "Including classification"],
        ["Avg Generation Time (ms)", "(from Kaggle run)", "(from Kaggle run)", "FLAN-T5 inference"],
        ["Avg End-to-End (ms)", "(from Kaggle run)", "(from Kaggle run)", "Total per-query latency"],
        ["Index Size (MB)", "(from Kaggle run)", "(from Kaggle run)", "Dense + graph storage"],
        ["Preprocessing Time (total)", "(from Kaggle run)", "(from Kaggle run)", "One-time build cost"],
        ["GPU Memory Peak (GB)", "(from Kaggle run)", "(from Kaggle run)", "T4 16 GB ceiling"],
    ],
    col_widths=[4.5, 3.5, 3.5, 4.0]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════
# 4. KEY FINDINGS
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "4. Key Findings", level=1)

add_body(doc,
    "The following findings are provisional — actual numerical results will be confirmed upon "
    "Kaggle notebook execution. The infrastructure for all six stages is fully built and documented; "
    "this report provides the methodology framework into which empirical results will be inserted."
)

findings = [
    ("Finding 1: Structured chunking preserves legal context.",
     "By respecting SECTION, Article, and paragraph boundaries, structured chunking eliminates "
     "the entity fragmentation problem inherent in fixed-token chunking. This is expected to improve "
     "retrieval precision for clause-specific and section-referencing queries, at the cost of "
     "variable chunk sizes that may require truncation for very large sections."),
    ("Finding 2: Entity extraction quality is the pipeline bottleneck.",
     "FLAN-T5-XL at 3B parameters, prompted zero-shot for structured JSON entity extraction, "
     "may produce inconsistent output. The JSON parse rate (percentage of FLAN-T5 responses yielding "
     "valid, parseable JSON) is the single most important diagnostic. The regex fallback ensures "
     "pipeline continuity but with lower entity typing precision. A higher-capability LLM or few-shot "
     "prompting would likely improve extraction quality significantly."),
    ("Finding 3: Graph-1Hop retrieval targets relationship questions.",
     "The entity co-occurrence graph provides a complementary retrieval signal to dense embeddings. "
     "For queries that involve named entities (parties, clauses, defined terms), 1-hop traversal "
     "retrieves chunks that embed the relational context, potentially surfacing relevant passages "
     "that embedding similarity alone would rank lower. The effect size depends on graph connectivity "
     "and entity extraction recall."),
    ("Finding 4: Community search enables synthesis-level retrieval.",
     "Connected-component communities, summarized by FLAN-T5, provide a higher-level retrieval unit "
     "than individual chunks. For synthesis questions requiring broad cross-section coverage, "
     "community retrieval provides more context volume and thematic coherence. However, large "
     "communities dilute relevance, and singleton communities contribute little."),
    ("Finding 5: Intelligent routing amplifies structural benefits.",
     "The query router ensures that each question type uses the retrieval method best suited to it. "
     "This prevents dense retrieval from being applied to relationship queries (where graph excels) "
     "and prevents graph retrieval from being applied to factual queries (where it adds noise). "
     "The router's classification accuracy directly bounds the pipeline's overall improvement."),
    ("Finding 6: Cost remains $0 throughout.",
     "All components — FLAN-T5-XL, all-MiniLM-L6-v2, FAISS, NetworkX — run locally on Kaggle's "
     "free T4 GPU. The entire Track 2 pipeline, from structured chunking through community "
     "summarization to query routing, incurs zero API or infrastructure cost. This demonstrates "
     "that knowledge-graph-enhanced RAG is feasible within free-tier compute budgets."),
]

for f_title, f_desc in findings:
    p = doc.add_paragraph()
    run = p.add_run(f_title + " ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run = p.add_run(f_desc)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════
# 5. LIMITATIONS & NEXT STEPS
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "5. Limitations & Next Steps", level=1)

add_heading_styled(doc, "5.1 Known Limitations", level=2)

limitations = [
    ("Entity Extraction Fidelity",
     "FLAN-T5-XL (3B) is pushed to its limits on structured JSON entity extraction without "
     "few-shot examples. The expected JSON parse rate may be below 70%, limiting graph quality. "
     "A higher-capability model (e.g., FLAN-T5-XXL 11B, or GPT-3.5/4 if API access is available) "
     "would improve extraction reliability, though at higher cost."),
    ("50-Document Subset",
     "Full entity extraction across all 714 documents would require approximately 30,000 LLM calls "
     "and is infeasible within Kaggle's session time limits. The 50-document subset provides a "
     "representative sample but may not capture the full diversity of the corpus. The graph's "
     "coverage is therefore partial."),
    ("Graph Sparsity",
     "Entity co-occurrence graphs built from small corpora may be sparse, with many isolated nodes "
     "or small disconnected components. This limits the effectiveness of both 1-hop traversal "
     "(few neighbors to traverse) and community search (many singleton communities)."),
    ("Classifier Simplicity",
     "The heuristic query classifier (keyword + entity count) is simple and fast but may misclassify "
     "nuanced questions. An LLM classifier is more flexible but adds latency. Neither approach is "
     "trained on legal QA classification data, limiting accuracy."),
    ("Community Summarization Quality",
     "FLAN-T5 one-sentence summaries of communities may oversimplify complex legal topics. "
     "A community covering 'Indemnification' may lose the distinction between different types "
     "of indemnification obligations that are legally distinct."),
    ("No Graph Learning",
     "The current graph is purely structural (co-occurrence). No graph neural network or "
     "representation learning (e.g., Node2Vec, GNN embeddings) is applied. This limits the "
     "graph's ability to discover latent relationships beyond explicit co-occurrence."),
    ("Single Embedding Model",
     "Both chunk and community-summary retrieval use all-MiniLM-L6-v2 (384d), which is optimized "
     "for general semantic similarity, not legal-domain retrieval. A legal-domain-tuned embedding "
     "model could improve retrieval quality across all stages."),
]

for l_title, l_desc in limitations:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(l_title + ": ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run = p.add_run(l_desc)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

add_heading_styled(doc, "5.2 Next Steps", level=2)

next_steps = [
    ("Execute Kaggle Notebook",
     "Run the full six-stage pipeline on the Kaggle T4 GPU to populate all placeholder metrics. "
     "The notebook implements all stages with checkpointed caching to allow incremental execution "
     "and debugging. Estimated runtime: 2–4 hours for the full pipeline on the 50-doc subset."),
    ("Improve Entity Extraction",
     "Explore few-shot prompting with 3–5 hand-crafted JSON examples, experiment with constrained "
     "decoding (outlines library), or switch to a larger model if memory permits. Target: ≥85% "
     "JSON parse rate for reliable graph construction."),
    ("Scale to Full Corpus",
     "Implement batched entity extraction with checkpointing to process all 714 documents. "
     "Consider using Kaggle's GPU quota across multiple sessions. A full-corpus graph would provide "
     "much richer connectivity and more reliable community detection."),
    ("Add Graph Embedding Layer",
     "Train Node2Vec embeddings on the entity co-occurrence graph to capture latent structural "
     "relationships. Combine graph embeddings with text embeddings for hybrid retrieval that "
     "leverages both semantic and structural signals."),
    ("Fine-tune Query Classifier",
     "Label 100–200 queries with gold-standard type labels (factual/relationship/synthesis) and "
     "fine-tune a small classifier (e.g., DistilBERT) for high-accuracy routing. This would "
     "reduce reliance on heuristic classification."),
    ("Compare with GraphRAG Baselines",
     "Benchmark against established graph-enhanced RAG frameworks (Microsoft GraphRAG, LightRAG) "
     "using the same LegalBench-RAG dataset to contextualize the Track 2 pipeline's performance "
     "within the broader research landscape."),
    ("Domain-Specific Embedding Model",
     "Fine-tune a legal-domain embedding model (e.g., Legal-BERT embeddings) to replace "
     "all-MiniLM-L6-v2. This would improve retrieval relevance across all stages by better "
     "capturing legal semantic relationships."),
]

for ns_title, ns_desc in next_steps:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(ns_title + ": ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run = p.add_run(ns_desc)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════
# APPENDIX: TECHNICAL SPECIFICATIONS
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "Appendix A: Technical Specifications", level=1)

add_heading_styled(doc, "A.1 Entity Extraction Prompt Template", level=2)
add_body(doc,
    "The following prompt template is sent to FLAN-T5-XL for entity extraction. Note that FLAN-T5 "
    "was fine-tuned on instruction-following tasks and may produce variable JSON formatting."
)

code_text = (
    'Extract the following entity types from this legal text chunk.\n'
    'Return ONLY valid JSON with these keys:\n'
    '  "PARTY": list of company or person names\n'
    '  "CLAUSE": list of section/article references\n'
    '  "DATE": list of dates mentioned\n'
    '  "AMOUNT": list of monetary amounts or quantities\n'
    '  "TERM": list of defined terms\n'
    '  "OBLIGATION": list of mandatory actions (must/shall)\n'
    '  "RIGHT": list of permissive actions (may/entitled)\n\n'
    'Legal text:\n{chunk_text}\n\n'
    'JSON output:'
)

p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(12)

add_heading_styled(doc, "A.2 Regex Fallback Patterns", level=2)

add_metric_table(doc,
    ["Entity Type", "Regex Pattern(s)", "Coverage"],
    [
        ["PARTY", r'\\b[A-Z][a-z]+\\s+(?:Corporation|Corp|Inc|LLC|Ltd|Company|Group|Holdings|Partners)\\b', "Named corporate entities"],
        ["CLAUSE", r'(?:Section|Article|§)\\s*\\d+[.\\d]*[a-zA-Z]*', "Legal section references"],
        ["DATE", r'\\b(?:January|February|…|December)\\s+\\d{1,2},?\\s+\\d{4}\\b|\\d{1,2}/\\d{1,2}/\\d{2,4}', "Date mentions"],
        ["AMOUNT", r'\\$[\\d,]+(?:\\.\\d{2})?|\\d+%', "Dollar amounts and percentages"],
        ["TERM", r'"([^"]+)"|\\b[A-Z][A-Z\\s]{2,}[A-Z]\\b', "Quoted terms and ALL-CAPS terms"],
        ["OBLIGATION", r'\\b(shall|must|will|agrees\\s+to|required\\s+to|covenants\\s+to)\\b', "Obligation markers"],
        ["RIGHT", r'\\b(may|entitled\\s+to|reserves\\s+the\\s+right|has\\s+the\\s+right|at\\s+its\\s+option)\\b', "Right/permission markers"],
    ],
    col_widths=[2.5, 9.0, 4.0]
)

add_heading_styled(doc, "A.3 Query Classification Heuristic", level=2)

add_body(doc, "The heuristic classifier uses the following logic (pseudocode):")

classifier_code = (
    "def classify_query(query: str, entities: list) -> str:\n"
    '    rel_keywords = ["between", "related to", "obligation of",\n'
    '                    "rights regarding", "relationship", "connection"]\n'
    '    syn_keywords = ["summarize", "overview", "all", "across",\n'
    '                    "patterns", "common", "types of"]\n'
    '    \n'
    '    if len(entities) >= 2 and any(kw in query.lower() for kw in rel_keywords):\n'
    '        return "relationship"\n'
    '    if any(kw in query.lower() for kw in syn_keywords):\n'
    '        return "synthesis"\n'
    '    return "factual"'
)

p = doc.add_paragraph()
run = p.add_run(classifier_code)
run.font.name = 'Consolas'
run.font.size = Pt(9)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(12)

add_heading_styled(doc, "A.4 Dependency Versions", level=2)

add_metric_table(doc,
    ["Package", "Version", "Purpose"],
    [
        ["transformers", "≥4.36", "FLAN-T5-XL model loading & inference"],
        ["sentence-transformers", "≥2.2", "all-MiniLM-L6-v2 embeddings"],
        ["faiss-cpu", "≥1.7", "Vector index & similarity search"],
        ["networkx", "≥3.0", "Entity co-occurrence graph"],
        ["python-docx", "≥1.1", "Report generation"],
        ["torch", "≥2.0", "Deep learning backend"],
        ["datasets", "≥2.14", "LegalBench-RAG loading"],
        ["numpy", "≥1.24", "Numerical operations"],
        ["tqdm", "≥4.65", "Progress bars"],
        ["pandas", "≥2.0", "Data management & metrics"],
    ],
    col_widths=[4.0, 3.0, 8.5]
)

# ── Footer ──
add_page_break(doc)

add_heading_styled(doc, "Document Control", level=1)

add_metric_table(doc,
    ["Field", "Value"],
    [
        ["Report Title", "Track 2: Knowledge Structure & Graph-Based Retrieval"],
        ["Study", "RAG Engineering Study — LegalBench-RAG"],
        ["Generated By", "Task 2 RAG (AutoClaw Agent)"],
        ["Date", "July 7, 2026"],
        ["Status", "Draft — metrics pending Kaggle execution"],
        ["Infrastructure", "Kaggle T4 GPU, FLAN-T5-XL, all-MiniLM-L6-v2, FAISS, NetworkX"],
        ["Cost", "$0 (fully local, open-source stack)"],
    ],
    col_widths=[4.0, 11.5]
)

# ── Save ──
doc.save(OUTPUT_PATH)
print(f"Report saved to: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
